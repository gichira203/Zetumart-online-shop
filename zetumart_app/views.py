from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth.decorators import login_required
from django.contrib.auth import authenticate, login, logout
from django.contrib.auth.models import User
from django.contrib import messages
from django.http import JsonResponse, HttpResponse
from django.views.decorators.csrf import csrf_exempt
from django.views.decorators.http import require_POST
from django.utils import timezone
from django.core.mail import send_mail
from django.db.models import Q
import json
import datetime
import openpyxl
from openpyxl.styles import Font, PatternFill
from docx import Document
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from .models import Product, Category, Order, ChatMessage, OrderTracking, AdminUser, UserProfile, CustomerCareMessage, Notification
from .payment_services import PaymentProcessor
from .delivery_data import get_county_list, search_counties, get_cities_for_county, search_cities, get_delivery_info, calculate_delivery_fee, get_estimated_delivery_date
import io

# Create your views here.

def index(request):
    unread_notifications = 0
    if request.user.is_authenticated:
        unread_notifications = Notification.objects.filter(user=request.user, is_read=False).count()
    
    # Get active products for display
    products = Product.objects.filter(status='active').select_related('category').order_by('-created_at')
    
    context = {
        'unread_notifications': unread_notifications,
        'products': products
    }
    return render(request, 'index.html', context)

def admin_dashboard(request):
    # Get all users with their profiles
    users = User.objects.all().select_related('userprofile').order_by('-date_joined')
    
    context = {
        'users': users,
    }
    return render(request, 'admin.html', context)

def login_view(request):
    if request.method == 'POST':
        username = request.POST.get('username')
        password = request.POST.get('password')
        
        # Check for admin credentials
        if username == 'admin' and password == 'root':
            # Create or get admin user
            user, created = User.objects.get_or_create(
                username='admin',
                defaults={'is_staff': True, 'is_superuser': True}
            )
            if created:
                user.set_password('root')
                user.save()
            
            # Authenticate and login
            user = authenticate(request, username='admin', password='root')
            if user:
                login(request, user)
                return redirect('admin')
        
        # Regular user authentication (for future use)
        user = authenticate(request, username=username, password=password)
        if user:
            login(request, user)
            return redirect('index')
        else:
            messages.error(request, 'Invalid credentials')
    
    return render(request, 'login.html')

def register_view(request):
    if request.method == 'POST':
        first_name = request.POST.get('firstName')
        last_name = request.POST.get('lastName')
        username = request.POST.get('username')
        email = request.POST.get('email')
        phone = request.POST.get('phone')
        password = request.POST.get('password')
        confirm_password = request.POST.get('confirmPassword')
        
        if password != confirm_password:
            messages.error(request, 'Passwords do not match')
            return render(request, 'register.html')
        
        # Create user (for future use)
        try:
            user = User.objects.create_user(
                username=username,
                email=email,
                password=password,
                first_name=first_name,
                last_name=last_name
            )
            
            # Create user profile
            UserProfile.objects.create(
                user=user,
                phone=phone
            )
            
            messages.success(request, 'Account created successfully!')
            return redirect('login')
        except Exception as e:
            messages.error(request, 'Username or email already exists')
    
    return render(request, 'register.html')

def contact_view(request):
    # Clear any existing messages to prevent showing old logout messages
    storage = messages.get_messages(request)
    storage.used = True
    
    if request.method == 'POST':
        name = request.POST.get('name')
        email = request.POST.get('email')
        phone = request.POST.get('phone', '')
        subject = request.POST.get('subject')
        message = request.POST.get('message')
        
        # Store message in database for admin to view
        try:
            CustomerCareMessage.objects.create(
                user=request.user if request.user.is_authenticated else None,
                name=name,
                email=email,
                phone=phone,
                message=message,
                message_type='agent',
                subject=subject
            )
            
            # Check if request is AJAX (from our JavaScript)
            if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
                return JsonResponse({'success': True, 'message': 'Message sent successfully!'})
            
            return redirect('contact')
            
        except Exception as e:
            # Check if request is AJAX (from our JavaScript)
            if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
                return JsonResponse({'success': False, 'error': 'There was an error sending your message. Please try again.'})
            
            messages.error(request, 'There was an error sending your message. Please try again.')
    
    return render(request, 'contact.html')

def logout_view(request):
    logout(request)
    # Clear any existing messages to avoid duplicates
    storage = messages.get_messages(request)
    storage.used = True
    messages.success(request, 'You have been logged out successfully')
    return redirect('index')

@login_required
def profile_view(request):
    profile, created = UserProfile.objects.get_or_create(user=request.user)
    unread_notifications = Notification.objects.filter(user=request.user, is_read=False).count()
    unread_messages = CustomerCareMessage.objects.filter(user=request.user, is_read=False).count()
    
    context = {
        'profile': profile,
        'unread_notifications': unread_notifications,
        'unread_messages': unread_messages
    }
    return render(request, 'profile.html', context)

@login_required
def update_profile(request):
    if request.method == 'POST':
        profile = UserProfile.objects.get(user=request.user)
        user = request.user
        
        # Update user info
        user.first_name = request.POST.get('first_name', user.first_name)
        user.last_name = request.POST.get('last_name', user.last_name)
        user.email = request.POST.get('email', user.email)
        
        # Update password if provided
        new_password = request.POST.get('new_password')
        if new_password:
            user.set_password(new_password)
        
        user.save()
        
        # Update profile
        profile.phone = request.POST.get('phone', profile.phone)
        profile.address = request.POST.get('address', profile.address)
        
        # Handle profile picture upload
        if 'profile_picture' in request.FILES:
            profile.profile_picture = request.FILES['profile_picture']
        
        profile.save()
        
        messages.success(request, 'Profile updated successfully!')
        return redirect('profile')
    
    return redirect('profile')

def customer_care_view(request):
    if request.method == 'POST':
        message_type = request.POST.get('message_type', 'agent')
        message = request.POST.get('message')
        
        if request.user.is_authenticated:
            # Logged-in user
            CustomerCareMessage.objects.create(
                user=request.user,
                message_type=message_type,
                message=message
            )
        else:
            # Non-logged-in user
            name = request.POST.get('name')
            email = request.POST.get('email')
            phone = request.POST.get('phone', '')
            whatsapp = request.POST.get('whatsapp', '')
            
            CustomerCareMessage.objects.create(
                user=None,
                name=name,
                email=email,
                phone=phone,
                whatsapp=whatsapp,
                message_type=message_type,
                message=message
            )
        
        if message_type == 'ai':
            messages.info(request, 'AI Chat feature coming soon!')
        else:
            messages.success(request, 'Message sent to agent successfully!')
        
        return redirect('customer_care')
    
    # Get messages
    if request.user.is_authenticated:
        messages_list = CustomerCareMessage.objects.filter(user=request.user).order_by('-created_at')
    else:
        messages_list = []
    
    context = {
        'messages': messages_list
    }
    return render(request, 'customer_care.html', context)

@login_required
def notifications_view(request):
    notifications_list = Notification.objects.filter(user=request.user).order_by('-created_at')
    
    # Mark notifications as read
    Notification.objects.filter(user=request.user, is_read=False).update(is_read=True)
    
    context = {
        'notifications': notifications_list
    }
    return render(request, 'notifications.html', context)

def admin_dashboard_view(request):
    if not request.user.is_staff:
        return redirect('index')
    
    # Create sample orders if none exist (for testing)
    if Order.objects.count() == 0:
        create_sample_orders()
    
    users_list = User.objects.all()
    messages_list = CustomerCareMessage.objects.all().order_by('-created_at')
    orders_list = Order.objects.all().order_by('-created_at')
    
    # Create default categories if they don't exist
    default_categories = ['Electronics', 'Fashion', 'Shoes', 'Home & Garden', 'Sports', 'Books', 'Toys', 'Beauty']
    for cat_name in default_categories:
        Category.objects.get_or_create(name=cat_name)
    
    categories_list = Category.objects.all()
    
    # Calculate statistics
    replied_count = messages_list.filter(is_replied=True).count()
    pending_count = messages_list.filter(is_replied=False).count()
    
    # Serialize messages for JavaScript
    messages_data = []
    for msg in messages_list:
        messages_data.append({
            'id': msg.id,
            'message': msg.message,
            'reply': msg.reply,
            'is_replied': msg.is_replied,
            'created_at': msg.created_at.isoformat(),
            'replied_at': msg.replied_at.isoformat() if msg.replied_at else None,
            'user': {
                'first_name': msg.user.first_name,
                'last_name': msg.user.last_name
            } if msg.user else None,
            'name': msg.name,
            'email': msg.email
        })
    
    # Serialize orders for JavaScript - FIX THIS PART
    orders_data = []
    for order in orders_list:
        orders_data.append({
            'id': order.id,
            'order_id': order.order_id,
            'customer_name': order.customer_name,
            'customer_email': order.customer_email,
            'customer_phone': order.customer_phone,
            'total_amount': float(order.total_amount),
            'payment_method': order.get_payment_method_display(),
            'payment_status': order.get_payment_status_display(),
            'order_status': order.get_order_status_display(),
            'delivery_status': order.get_delivery_status_display(),
            'created_at': order.created_at.strftime('%Y-%m-%d %H:%M'),
            'delivery_county': order.delivery_county,
            'delivery_town': order.delivery_town
        })
    
    context = {
        'users': users_list,
        'messages': messages_list,
        'messages_json': json.dumps(messages_data),
        'orders': orders_list,
        'orders_json': json.dumps(orders_data),
        'categories': categories_list,
        'products': Product.objects.all().select_related('category').order_by('-created_at'),
        'replied_count': replied_count,
        'pending_count': pending_count,
    }
    return render(request, 'admin_dashboard.html', context)

def create_sample_orders():
    """Create sample orders for testing the dashboard"""
    import random
    from datetime import datetime, timedelta
    
    sample_customers = [
        {'name': 'John Doe', 'email': 'john@example.com', 'phone': '0712345678'},
        {'name': 'Jane Smith', 'email': 'jane@example.com', 'phone': '0723456789'},
        {'name': 'Bob Johnson', 'email': 'bob@example.com', 'phone': '0734567890'},
        {'name': 'Alice Brown', 'email': 'alice@example.com', 'phone': '0745678901'},
        {'name': 'Charlie Wilson', 'email': 'charlie@example.com', 'phone': '0756789012'}
    ]
    
    counties = ['Nairobi', 'Mombasa', 'Kisumu', 'Nakuru', 'Eldoret']
    towns = ['Nairobi CBD', 'Mombasa Town', 'Kisumu Central', 'Nakuru Town', 'Eldoret Town']
    statuses = ['pending', 'confirmed', 'processing', 'shipped', 'delivered', 'cancelled']
    payment_methods = ['mpesa', 'till', 'cod', 'card']
    payment_statuses = ['pending', 'completed', 'failed']
    
    for i in range(10):
        customer = random.choice(sample_customers)
        county = random.choice(counties)
        town = random.choice(towns)
        status = random.choice(statuses)
        payment_method = random.choice(payment_methods)
        payment_status = random.choice(payment_statuses)
        
        # Generate random date within last 30 days
        days_ago = random.randint(0, 30)
        created_date = timezone.now() - timedelta(days=days_ago)
        
        Order.objects.create(
            order_id=f'ORD-{1000 + i}',
            customer_name=customer['name'],
            customer_email=customer['email'],
            customer_phone=customer['phone'],
            items=[{'name': f'Product {i+1}', 'quantity': random.randint(1, 3), 'price': random.uniform(1000, 10000)}],
            subtotal=random.uniform(1000, 10000),
            delivery_fee=random.uniform(100, 500),
            total_amount=random.uniform(1100, 10500),
            delivery_county=county,
            delivery_town=town,
            delivery_point=f'{town}, {county}',
            delivery_method=random.choice(['home', 'office', 'pickup']),
            delivery_time=random.choice(['anytime', 'morning', 'afternoon', 'evening']),
            payment_method=payment_method,
            payment_status=payment_status,
            order_status=status,
            delivery_status=random.choice(['pending', 'preparing', 'out_for_delivery', 'delivered']),
            created_at=created_date
        )

@csrf_exempt
@login_required
def reply_to_message(request, message_id):
    if not request.user.is_staff:
        return JsonResponse({'error': 'Unauthorized'}, status=403)
    
    message = get_object_or_404(CustomerCareMessage, id=message_id)
    
    if request.method == 'POST':
        reply = request.POST.get('reply')
        message.reply = reply
        message.is_replied = True
        message.replied_at = timezone.now()
        message.save()
        
        # Create notification for user
        if message.user:
            Notification.objects.create(
                user=message.user,
                message=f"Admin replied to your customer care message"
            )
        
        # Send email notification
        try:
            subject = 'ZetuMart - Response to Your Customer Care Message'
            email_message = f"""
            Dear {message.user.first_name if message.user else message.name or 'Customer'},
            
            Thank you for contacting ZetuMart Customer Care. We have responded to your message:
            
            Your original message:
            "{message.message}"
            
            Our response:
            "{reply}"
            
            You can view this response in your account dashboard under notifications.
            
            Best regards,
            ZetuMart Customer Care Team
            """
            
            recipient_emails = []
            if message.user and message.user.email:
                recipient_emails.append(message.user.email)
            if message.email:
                recipient_emails.append(message.email)
            
            if recipient_emails:
                send_mail(
                    subject,
                    email_message,
                    'noreply@zetumart.com',
                    recipient_emails,
                    fail_silently=False,
                )
        except Exception as e:
            print(f"Email sending failed: {e}")
        
        # Send WhatsApp notification (you would need to integrate with WhatsApp API)
        if message.whatsapp:
            try:
                # This is a placeholder for WhatsApp API integration
                # You would need to use services like Twilio WhatsApp API
                whatsapp_message = f"ZetuMart Response: {reply[:100]}..."
                print(f"WhatsApp notification would be sent to {message.whatsapp}: {whatsapp_message}")
                # Actual implementation would require WhatsApp Business API
            except Exception as e:
                print(f"WhatsApp sending failed: {e}")
        
        return JsonResponse({'success': True})
    
    return JsonResponse({'error': 'Invalid request'}, status=400)

def export_users_excel(request):
    if not request.user.is_staff:
        return redirect('login')
    
    users = User.objects.all()
    
    # Create Excel workbook
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Users Report"
    
    # Define headers based on actual UserProfile fields
    headers = [
        'ID', 'Username', 'Email', 'First Name', 'Last Name', 
        'Phone', 'Date Joined', 'Last Login', 'Is Active', 'Is Staff',
        'Profile Picture', 'Address', 'Date of Birth'
    ]
    
    # Apply headers
    for col_num, header in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col_num)
        cell.value = header
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = PatternFill(start_color="366092", end_color="366092", fill_type="solid")
    
    # Add user data
    for row_num, user in enumerate(users, 2):
        try:
            profile = user.userprofile
        except UserProfile.DoesNotExist:
            profile = None
            
        ws.cell(row=row_num, column=1, value=user.id)
        ws.cell(row=row_num, column=2, value=user.username)
        ws.cell(row=row_num, column=3, value=user.email)
        ws.cell(row=row_num, column=4, value=user.first_name or '')
        ws.cell(row=row_num, column=5, value=user.last_name or '')
        ws.cell(row=row_num, column=6, value=profile.phone if profile else '')
        ws.cell(row=row_num, column=7, value=user.date_joined.strftime('%Y-%m-%d %H:%M:%S'))
        ws.cell(row=row_num, column=8, value=user.last_login.strftime('%Y-%m-%d %H:%M:%S') if user.last_login else 'Never')
        ws.cell(row=row_num, column=9, value='Yes' if user.is_active else 'No')
        ws.cell(row=row_num, column=10, value='Yes' if user.is_staff else 'No')
        ws.cell(row=row_num, column=11, value='Yes' if profile and profile.profile_picture else 'No')
        ws.cell(row=row_num, column=12, value=profile.address if profile else '')
        ws.cell(row=row_num, column=13, value=profile.date_of_birth.strftime('%Y-%m-%d') if profile and profile.date_of_birth else '')
    
    # Adjust column widths
    for column in ws.columns:
        max_length = 0
        column_letter = column[0].column_letter
        for cell in column:
            try:
                if len(str(cell.value)) > max_length:
                    max_length = len(str(cell.value))
            except:
                pass
        adjusted_width = min(max_length + 2, 50)
        ws.column_dimensions[column_letter].width = adjusted_width
    
    # Create response
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    response['Content-Disposition'] = 'attachment; filename="users_report.xlsx"'
    
    # Save workbook to response
    wb.save(response)
    
    return response

def export_users_word(request):
    if not request.user.is_staff:
        return redirect('login')
    
    users = User.objects.all()
    
    doc = Document()
    doc.add_heading('ZetuMart Users Data', 0)
    
    # Add summary
    doc.add_paragraph(f'Total Users: {users.count()}')
    doc.add_paragraph(f'Generated on: {timezone.now().strftime("%Y-%m-%d %H:%M:%S")}')
    doc.add_paragraph('')
    
    # Create table with actual UserProfile fields
    table = doc.add_table(rows=1, cols=13)
    table.style = 'Table Grid'
    
    # Define headers based on actual UserProfile fields
    headers = [
        'ID', 'Username', 'Email', 'First Name', 'Last Name', 
        'Phone', 'Date Joined', 'Last Login', 'Is Active', 'Is Staff',
        'Profile Picture', 'Address', 'Date of Birth'
    ]
    
    # Add headers to table
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        # Make header bold
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    
    # Add user data to table
    for user in users:
        try:
            profile = user.userprofile
        except UserProfile.DoesNotExist:
            profile = None
        
        row_cells = table.add_row().cells
        row_cells[0].text = str(user.id)
        row_cells[1].text = user.username or ''
        row_cells[2].text = user.email or ''
        row_cells[3].text = user.first_name or ''
        row_cells[4].text = user.last_name or ''
        row_cells[5].text = str(profile.phone) if profile and profile.phone else ''
        row_cells[6].text = user.date_joined.strftime('%Y-%m-%d %H:%M:%S')
        row_cells[7].text = user.last_login.strftime('%Y-%m-%d %H:%M:%S') if user.last_login else 'Never'
        row_cells[8].text = 'Yes' if user.is_active else 'No'
        row_cells[9].text = 'Yes' if user.is_staff else 'No'
        row_cells[10].text = 'Yes' if profile and profile.profile_picture else 'No'
        row_cells[11].text = profile.address if profile and profile.address else ''
        row_cells[12].text = profile.date_of_birth.strftime('%Y-%m-%d') if profile and profile.date_of_birth else ''
    
    # Add summary statistics section
    doc.add_page_break()
    doc.add_heading('Summary Statistics', level=1)
    
    active_users = users.filter(is_active=True).count()
    staff_users = users.filter(is_staff=True).count()
    users_with_profiles = UserProfile.objects.all().count()
    users_with_pictures = UserProfile.objects.exclude(profile_picture='profile_pics/default.jpg').count()
    
    stats_paragraph = doc.add_paragraph()
    stats_paragraph.add_run('Active Users: ').bold = True
    stats_paragraph.add_run(f'{active_users} / {users.count()}\n')
    stats_paragraph.add_run('Staff Users: ').bold = True
    stats_paragraph.add_run(f'{staff_users}\n')
    stats_paragraph.add_run('Users with Complete Profiles: ').bold = True
    stats_paragraph.add_run(f'{users_with_profiles}\n')
    stats_paragraph.add_run('Users with Profile Pictures: ').bold = True
    stats_paragraph.add_run(f'{users_with_pictures}\n')
    
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename=users_data.docx'
    doc.save(response)
    return response

@login_required
def admin_user_management(request):
    if not request.user.is_superuser:
        return JsonResponse({'error': 'Superuser access required'}, status=403)
    
    if request.method == 'GET':
        users = User.objects.all().prefetch_related('userprofile')
        users_data = []
        
        for user in users:
            try:
                profile = user.userprofile
            except UserProfile.DoesNotExist:
                profile = None
                
            users_data.append({
                'id': user.id,
                'username': user.username,
                'email': user.email,
                'first_name': user.first_name,
                'last_name': user.last_name,
                'is_active': user.is_active,
                'is_staff': user.is_staff,
                'is_superuser': user.is_superuser,
                'date_joined': user.date_joined.strftime('%Y-%m-%d %H:%M:%S'),
                'last_login': user.last_login.strftime('%Y-%m-%d %H:%M:%S') if user.last_login else 'Never',
                'phone': profile.phone if profile else '',
                'address': profile.address if profile else '',
                'profile_picture': profile.profile_picture.url if profile and profile.profile_picture else None,
                'password_hash': user.password,  # For superuser viewing only
            })
        
        return JsonResponse({'users': users_data})
    
    elif request.method == 'POST':
        action = request.POST.get('action')
        
        if action == 'delete_user':
            user_id = request.POST.get('user_id')
            user = get_object_or_404(User, id=user_id)
            if user.is_superuser:
                return JsonResponse({'error': 'Cannot delete superuser'}, status=400)
            user.delete()
            return JsonResponse({'success': True, 'message': 'User deleted successfully'})
        
        elif action == 'toggle_status':
            user_id = request.POST.get('user_id')
            user = get_object_or_404(User, id=user_id)
            user.is_active = not user.is_active
            user.save()
            status = 'activated' if user.is_active else 'deactivated'
            return JsonResponse({'success': True, 'message': f'User {status} successfully'})
        
        elif action == 'make_staff':
            user_id = request.POST.get('user_id')
            user = get_object_or_404(User, id=user_id)
            user.is_staff = not user.is_staff
            user.save()
            status = 'grprogrammed as staff' if user.is_staff else 'removed from staff'
            return JsonResponse({'success': True, 'message': f'User {status} successfully'})
        
        elif action == 'reset_password':
            user_id = request.POST.get('user_id')
            new_password = request.POST.get('new_password')
            user = get_object_or_404(User, id=user_id)
            user.set_password(new_password)
            user.save()
            return JsonResponse({'success': True, 'message': 'Password reset successfully'})
    
    return JsonResponse({'error': 'Invalid request'}, status=400)

@login_required
def admin_product_management(request):
    if not request.user.is_superuser:
        return JsonResponse({'error': 'Superuser access required'}, status=403)
    
    if request.method == 'POST':
        action = request.POST.get('action')
        product_id = request.POST.get('product_id')
        
        if action == 'delete_product':
            product = get_object_or_404(Product, id=product_id)
            product.delete()
            return JsonResponse({'success': True, 'message': 'Product deleted successfully'})
        
        elif action == 'toggle_status':
            product = get_object_or_404(Product, id=product_id)
            product.status = 'inactive' if product.status == 'active' else 'active'
            product.save()
            status = 'activated' if product.status == 'active' else 'deactivated'
            return JsonResponse({'success': True, 'message': f'Product {status} successfully'})
    
    return JsonResponse({'error': 'Invalid request'}, status=400)

@login_required
def admin_message_management(request):
    if not request.user.is_superuser:
        return JsonResponse({'error': 'Superuser access required'}, status=403)
    
    if request.method == 'POST':
        action = request.POST.get('action')
        message_id = request.POST.get('message_id')
        
        if action == 'delete_message':
            message = get_object_or_404(CustomerCareMessage, id=message_id)
            message.delete()
            return JsonResponse({'success': True, 'message': 'Message deleted successfully'})
        
        elif action == 'mark_resolved':
            message = get_object_or_404(CustomerCareMessage, id=message_id)
            message.is_replied = True
    response = HttpResponse(content_type='application/pdf')
    response['Content-Disposition'] = 'attachment; filename=users_data.pdf'
    
    p = canvas.Canvas(response)
    p.setFont("Helvetica", 16)
    p.drawString(100, 800, "ZetuMart Users Data")
    
    y_position = 750
    p.setFont("Helvetica", 12)
    
    for user in users:
        if y_position < 100:
            p.showPage()
            y_position = 750
            
        try:
            profile = UserProfile.objects.get(user=user)
            user_data = [
                f"Username: {user.username}",
                f"Name: {user.first_name} {user.last_name}",
                f"Email: {user.email}",
                f"Joined: {user.date_joined.strftime('%Y-%m-%d')}",
                f"Phone: {profile.phone or 'Not provided'}",
                f"Address: {profile.address or 'Not provided'}"
            ]
        except UserProfile.DoesNotExist:
            user_data = [
                f"Username: {user.username}",
                f"Name: {user.first_name} {user.last_name}",
                f"Email: {user.email}",
                f"Joined: {user.date_joined.strftime('%Y-%m-%d')}",
                f"Phone: Not provided",
                f"Address: Not provided"
            ]
        
        for data in user_data:
            p.drawString(100, y_position, data)
            y_position -= 20
        
        y_position -= 10
    
    p.save()
    return response

@csrf_exempt
@login_required
def admin_message_management(request):
    if not request.user.is_superuser:
        return JsonResponse({'error': 'Superuser access required'}, status=403)
    
    if request.method == 'POST':
        action = request.POST.get('action')
        message_id = request.POST.get('message_id')
        
        if action == 'delete_message':
            message = get_object_or_404(CustomerCareMessage, id=message_id)
            message.delete()
            return JsonResponse({'success': True, 'message': 'Message deleted successfully'})
        
        elif action == 'mark_resolved':
            message = get_object_or_404(CustomerCareMessage, id=message_id)
            message.is_replied = True
            message.replied_at = timezone.now()
            message.save()
            return JsonResponse({'success': True, 'message': 'Message marked as resolved'})
        
        elif action == 'bulk_delete':
            message_ids = request.POST.getlist('message_ids[]')
            CustomerCareMessage.objects.filter(id__in=message_ids).delete()
            return JsonResponse({'success': True, 'message': f'{len(message_ids)} messages deleted successfully'})
    
    return JsonResponse({'error': 'Invalid request'}, status=400)

def export_users_pdf(request):
    if not request.user.is_staff:
        return redirect('login')
    
    users = User.objects.all()
    
    response = HttpResponse(content_type='application/pdf')
    response['Content-Disposition'] = 'attachment; filename=users_data.pdf'
    
    p = canvas.Canvas(response)
    p.setFont("Helvetica", 16)
    p.drawString(100, 800, "ZetuMart Users Data")
    
    y_position = 750
    p.setFont("Helvetica", 12)
    
    for user in users:
        if y_position < 100:
            p.showPage()
            y_position = 750
            
        try:
            profile = UserProfile.objects.get(user=user)
            user_data = [
                f"Username: {user.username}",
                f"Name: {user.first_name} {user.last_name}",
                f"Email: {user.email}",
                f"Joined: {user.date_joined.strftime('%Y-%m-%d')}",
                f"Phone: {profile.phone or 'Not provided'}",
                f"Address: {profile.address or 'Not provided'}"
            ]
        except UserProfile.DoesNotExist:
            user_data = [
                f"Username: {user.username}",
                f"Name: {user.first_name} {user.last_name}",
                f"Email: {user.email}",
                f"Joined: {user.date_joined.strftime('%Y-%m-%d')}",
                f"Phone: Not provided",
                f"Address: Not provided"
            ]
        
        for data in user_data:
            p.drawString(100, y_position, data)
            y_position -= 20
        
        y_position -= 10
    
    p.save()
    return response

@csrf_exempt
@login_required
def add_admin_user(request):
    if not request.user.is_superuser:
        return JsonResponse({'error': 'Superuser access required'}, status=403)
    
    if request.method == 'POST':
        username = request.POST.get('username')
        email = request.POST.get('email')
        first_name = request.POST.get('first_name', '')
        last_name = request.POST.get('last_name', '')
        password = request.POST.get('password')
        
        if not username or not email or not password:
            return JsonResponse({'error': 'Username, email, and password are required'}, status=400)
        
        try:
            user = User.objects.create_user(
                username=username,
                email=email,
                password=password,
                first_name=first_name,
                last_name=last_name,
                is_staff=True
            )
            return JsonResponse({'success': True, 'message': 'Admin user created successfully'})
        except Exception as e:
            return JsonResponse({'error': str(e)}, status=400)
    
    return JsonResponse({'error': 'Invalid request method'}, status=405)

@csrf_exempt
@login_required
def delete_user(request, user_id):
    if not request.user.is_staff:
        return JsonResponse({'error': 'Unauthorized'}, status=403)
    
    user = get_object_or_404(User, id=user_id)
    
    if request.method == 'DELETE':
        if user.is_superuser:
            return JsonResponse({'error': 'Cannot delete superuser'}, status=400)
        user.delete()
        return JsonResponse({'success': True, 'message': 'User deleted successfully'})
    
    return JsonResponse({'error': 'Invalid request'}, status=400)

@csrf_exempt
@login_required
def delete_message(request, message_id):
    if not request.user.is_staff:
        return JsonResponse({'error': 'Unauthorized'}, status=403)
    
    message = get_object_or_404(CustomerCareMessage, id=message_id)
    
    if request.method == 'DELETE':
        message.delete()
        return JsonResponse({'success': True, 'message': 'Message deleted successfully'})
    
    return JsonResponse({'error': 'Invalid request'}, status=400)

@csrf_exempt
@login_required
def add_product(request):
    if request.method != 'POST':
        return JsonResponse({'error': 'Invalid request method'}, status=405)
    
    if not request.user.is_staff:
        return JsonResponse({'error': 'Unauthorized'}, status=403)
    
    try:
        # Debug - print received data
        print(f"POST data: {request.POST}")
        print(f"FILES data: {request.FILES}")
        print(f"User: {request.user}, is_staff: {request.user.is_staff}")
        
        # Get form data
        name = request.POST.get('name', '').strip()
        price = request.POST.get('price', '').strip()
        category_id = request.POST.get('category', '').strip()
        stock = request.POST.get('stock', '0').strip()
        description = request.POST.get('description', '').strip()
        status = request.POST.get('status', 'active')
        
        print(f"Processed data - name: '{name}', price: '{price}', category: '{category_id}', stock: '{stock}'")
        
        # Validate required fields
        if not name:
            print("ERROR: Product name is missing")
            return JsonResponse({'error': 'Product name is required'}, status=400)
        
        if not price:
            print("ERROR: Product price is missing")
            return JsonResponse({'error': 'Product price is required'}, status=400)
        
        try:
            price = float(price)
            if price <= 0:
                print(f"ERROR: Price {price} is not greater than 0")
                return JsonResponse({'error': 'Price must be greater than 0'}, status=400)
            # Removed upper limit - allow very high prices
        except ValueError as e:
            print(f"ERROR: Invalid price format: {e}")
            return JsonResponse({'error': 'Invalid price format'}, status=400)
        
        try:
            stock = int(stock)
            if stock < 0:
                print(f"ERROR: Stock {stock} is negative")
                return JsonResponse({'error': 'Stock cannot be negative'}, status=400)
            # Removed upper limit - allow very large stock quantities
        except ValueError as e:
            print(f"ERROR: Invalid stock format: {e}")
            return JsonResponse({'error': 'Invalid stock format'}, status=400)
        
        # Validate status
        valid_statuses = ['active', 'inactive']
        if status not in valid_statuses:
            status = 'active'  # Default to active if invalid
        
        # Get category (optional)
        category = None
        if category_id:
            try:
                category = Category.objects.get(id=int(category_id))
            except (Category.DoesNotExist, ValueError) as e:
                print(f"ERROR: Invalid category: {e}")
                return JsonResponse({'error': 'Invalid category'}, status=400)
        
        print("Creating product...")
        # Create product
        product = Product.objects.create(
            name=name,
            price=price,
            category=category,
            stock=stock,
            description=description,
            status=status
        )
        
        # Handle images
        for i in range(1, 5):
            image_field = f'image{i}'
            if image_field in request.FILES:
                setattr(product, image_field, request.FILES[image_field])
        
        product.save()
        print(f"Product created successfully: {product.id}")
        return JsonResponse({'success': True, 'message': 'Product added successfully'})
    
    except Exception as e:
        print(f"ERROR: Exception occurred: {e}")
        return JsonResponse({'error': f'Error adding product: {str(e)}'}, status=400)

@csrf_exempt
@login_required
def get_products(request):
    if not request.user.is_staff:
        return JsonResponse({'error': 'Unauthorized'}, status=403)
    
    products = Product.objects.all().select_related('category')
    products_data = []
    
    for product in products:
        products_data.append({
            'id': product.id,
            'name': product.name,
            'price': str(product.price),
            'category': product.category.name if product.category else 'No Category',
            'category_id': product.category.id if product.category else None,
            'stock': product.stock,
            'status': product.status,
            'description': product.description[:100] + '...' if product.description and len(product.description) > 100 else product.description,
            'image1': product.image1.url if product.image1 else None,
            'created_at': product.created_at.strftime('%Y-%m-%d %H:%M:%S')
        })
    
    return JsonResponse({'products': products_data})

@csrf_exempt
@login_required
def get_product(request, product_id):
    if not request.user.is_staff:
        return JsonResponse({'error': 'Unauthorized'}, status=403)
    
    try:
        product = Product.objects.get(id=product_id)
        return JsonResponse({
            'success': True,
            'product': {
                'id': product.id,
                'name': product.name,
                'price': str(product.price),
                'category_id': product.category.id if product.category else None,
                'stock': product.stock,
                'status': product.status,
                'description': product.description
            }
        })
    except Product.DoesNotExist:
        return JsonResponse({'error': 'Product not found'}, status=404)

@csrf_exempt
@login_required
def update_product(request, product_id):
    if not request.user.is_staff:
        return JsonResponse({'error': 'Unauthorized'}, status=403)
    
    if request.method != 'PUT':
        return JsonResponse({'error': 'Invalid request method'}, status=405)
    
    try:
        product = Product.objects.get(id=product_id)
        
        # Update product fields
        product.name = request.POST.get('name', product.name)
        product.price = request.POST.get('price', product.price)
        product.stock = request.POST.get('stock', product.stock)
        product.description = request.POST.get('description', product.description)
        product.status = request.POST.get('status', product.status)
        
        # Update category if provided
        category_id = request.POST.get('category')
        if category_id:
            category = Category.objects.get(id=category_id)
            product.category = category
        
        # Handle images
        for i in range(1, 5):
            image_field = f'image{i}'
            if image_field in request.FILES:
                setattr(product, image_field, request.FILES[image_field])
        
        product.save()
        return JsonResponse({'success': True, 'message': 'Product updated successfully'})
    
    except Product.DoesNotExist:
        return JsonResponse({'error': 'Product not found'}, status=404)
    except Category.DoesNotExist:
        return JsonResponse({'error': 'Invalid category'}, status=400)
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=400)

@csrf_exempt
@login_required
def delete_product_view(request, product_id):
    if not request.user.is_staff:
        return JsonResponse({'error': 'Unauthorized'}, status=403)
    
    if request.method != 'DELETE':
        return JsonResponse({'error': 'Invalid request method'}, status=405)
    
    try:
        product = Product.objects.get(id=product_id)
        product.delete()
        return JsonResponse({'success': True, 'message': 'Product deleted successfully'})
    except Product.DoesNotExist:
        return JsonResponse({'error': 'Product not found'}, status=404)

def products_view(request):
    products = Product.objects.filter(status='active').select_related('category')
    categories = Category.objects.all()
    
    context = {
        'products': products,
        'categories': categories
    }
    return render(request, 'products.html', context)

def shop_category_view(request):
    products = Product.objects.filter(status='active').select_related('category')
    categories = Category.objects.all()
    
    category_id = request.GET.get('category')
    if category_id:
        products = products.filter(category_id=category_id)
    
    context = {
        'products': products,
        'categories': categories,
        'selected_category': category_id
    }
    return render(request, 'shop.html', context)

@csrf_exempt
def get_categories(request):
    categories = Category.objects.all()
    categories_data = []
    
    for category in categories:
        categories_data.append({
            'id': category.id,
            'name': category.name
        })
    
    return JsonResponse({'categories': categories_data})

@csrf_exempt
@login_required
def create_order(request):
    if request.method != 'POST':
        return JsonResponse({'error': 'Invalid request method'}, status=405)
    
    try:
        data = json.loads(request.body)
        
        # Create order
        order = Order.objects.create(
            customer_name=data.get('customer_name'),
            customer_email=data.get('customer_email'),
            customer_phone=data.get('customer_phone'),
            delivery_county=data.get('delivery_county'),
            delivery_town=data.get('delivery_town'),
            delivery_address=data.get('delivery_address'),
            total_amount=data.get('total_amount'),
            payment_method=data.get('payment_method'),
            order_status='pending',
            payment_status='pending'
        )
        
        return JsonResponse({
            'success': True,
            'order_id': order.order_id,
            'message': 'Order created successfully'
        })
    
    except json.JSONDecodeError:
        return JsonResponse({'error': 'Invalid JSON'}, status=400)
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=400)

@csrf_exempt
@login_required
def get_orders(request):
    if not request.user.is_staff:
        return JsonResponse({'error': 'Unauthorized'}, status=403)
    
    orders = Order.objects.all().order_by('-created_at')
    orders_data = []
    
    for order in orders:
        orders_data.append({
            'id': order.id,
            'order_id': order.order_id,
            'customer_name': order.customer_name,
            'customer_email': order.customer_email,
            'customer_phone': order.customer_phone,
            'total_amount': float(order.total_amount),
            'payment_method': order.get_payment_method_display(),
            'payment_status': order.get_payment_status_display(),
            'order_status': order.get_order_status_display(),
            'delivery_status': order.get_delivery_status_display(),
            'created_at': order.created_at.strftime('%Y-%m-%d %H:%M'),
            'delivery_county': order.delivery_county,
            'delivery_town': order.delivery_town
        })
    
    return JsonResponse({'orders': orders_data})

@csrf_exempt
@login_required
def update_order_status(request, order_id):
    if not request.user.is_staff:
        return JsonResponse({'error': 'Unauthorized'}, status=403)
    
    if request.method != 'PUT':
        return JsonResponse({'error': 'Invalid request method'}, status=405)
    
    try:
        order = get_object_or_404(Order, id=order_id)
        data = json.loads(request.body)
        
        # Update order status
        if 'order_status' in data:
            order.order_status = data['order_status']
        
        if 'payment_status' in data:
            order.payment_status = data['payment_status']
        
        if 'delivery_status' in data:
            order.delivery_status = data['delivery_status']
        
        order.save()
        return JsonResponse({'success': True, 'message': 'Order status updated successfully'})
    
    except json.JSONDecodeError:
        return JsonResponse({'error': 'Invalid JSON'}, status=400)
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=400)

@csrf_exempt
@login_required
def get_order_tracking(request, order_id):
    if not request.user.is_staff:
        return JsonResponse({'error': 'Unauthorized'}, status=403)
    
    try:
        order = get_object_or_404(Order, id=order_id)
        
        tracking_data = {
            'order_id': order.order_id,
            'order_status': order.get_order_status_display(),
            'payment_status': order.get_payment_status_display(),
            'delivery_status': order.get_delivery_status_display(),
            'created_at': order.created_at.strftime('%Y-%m-%d %H:%M:%S'),
            'estimated_delivery': order.estimated_delivery_date.strftime('%Y-%m-%d') if order.estimated_delivery_date else None,
            'delivery_county': order.delivery_county,
            'delivery_town': order.delivery_town,
            'delivery_address': order.delivery_address
        }
        
        return JsonResponse({'success': True, 'tracking': tracking_data})
    
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=400)

@csrf_exempt
@login_required
def send_chat_message(request):
    if request.method != 'POST':
        return JsonResponse({'error': 'Invalid request method'}, status=405)
    
    try:
        data = json.loads(request.body)
        message = data.get('message')
        
        if not message:
            return JsonResponse({'error': 'Message is required'}, status=400)
        
        # Create chat message (placeholder implementation)
        chat_message = {
            'id': 1,
            'message': message,
            'sender': request.user.username,
            'timestamp': timezone.now().strftime('%Y-%m-%d %H:%M:%S')
        }
        
        return JsonResponse({'success': True, 'message': chat_message})
    
    except json.JSONDecodeError:
        return JsonResponse({'error': 'Invalid JSON'}, status=400)
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=400)

@csrf_exempt
@login_required
def get_chat_messages(request):
    # Placeholder implementation - returns sample messages
    messages = [
        {
            'id': 1,
            'message': 'Welcome to ZetuMart support!',
            'sender': 'support',
            'timestamp': timezone.now().strftime('%Y-%m-%d %H:%M:%S')
        }
    ]
    
    return JsonResponse({'success': True, 'messages': messages})

@csrf_exempt
@login_required
def reply_to_chat_message(request, message_id):
    if request.method != 'POST':
        return JsonResponse({'error': 'Invalid request method'}, status=405)
    
    try:
        data = json.loads(request.body)
        reply = data.get('reply')
        
        if not reply:
            return JsonResponse({'error': 'Reply is required'}, status=400)
        
        # Create reply message (placeholder implementation)
        reply_message = {
            'id': message_id + 1000,
            'message': reply,
            'sender': request.user.username,
            'timestamp': timezone.now().strftime('%Y-%m-%d %H:%M:%S'),
            'reply_to': message_id
        }
        
        return JsonResponse({'success': True, 'reply': reply_message})
    
    except json.JSONDecodeError:
        return JsonResponse({'error': 'Invalid JSON'}, status=400)
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=400)

def checkout_view(request):
    if request.method == 'POST':
        # Handle checkout form submission
        customer_name = request.POST.get('customer_name')
        customer_email = request.POST.get('customer_email')
        customer_phone = request.POST.get('customer_phone')
        delivery_county = request.POST.get('delivery_county')
        delivery_town = request.POST.get('delivery_town')
        delivery_address = request.POST.get('delivery_address')
        payment_method = request.POST.get('payment_method')
        total_amount = request.POST.get('total_amount')
        
        try:
            order = Order.objects.create(
                customer_name=customer_name,
                customer_email=customer_email,
                customer_phone=customer_phone,
                delivery_county=delivery_county,
                delivery_town=delivery_town,
                delivery_address=delivery_address,
                payment_method=payment_method,
                total_amount=total_amount,
                order_status='pending',
                payment_status='pending'
            )
            
            return JsonResponse({
                'success': True,
                'order_id': order.order_id,
                'message': 'Order placed successfully!'
            })
        except Exception as e:
            return JsonResponse({'error': str(e)}, status=400)
    
    # GET request - show checkout page
    return render(request, 'checkout.html')

def order_confirmation_view(request, order_id):
    try:
        order = get_object_or_404(Order, order_id=order_id)
        context = {
            'order': order
        }
        return render(request, 'order_confirmation.html', context)
    except Exception as e:
        return render(request, 'error.html', {'error': 'Order not found'})

@csrf_exempt
@login_required
def get_order_details(request, order_id):
    if not request.user.is_staff:
        return JsonResponse({'error': 'Unauthorized'}, status=403)
    
    try:
        order = get_object_or_404(Order, id=order_id)
        
        order_data = {
            'id': order.id,
            'order_id': order.order_id,
            'customer_name': order.customer_name,
            'customer_email': order.customer_email,
            'customer_phone': order.customer_phone,
            'delivery_county': order.delivery_county,
            'delivery_town': order.delivery_town,
            'delivery_address': order.delivery_address,
            'total_amount': float(order.total_amount),
            'payment_method': order.get_payment_method_display(),
            'payment_status': order.get_payment_status_display(),
            'order_status': order.get_order_status_display(),
            'delivery_status': order.get_delivery_status_display(),
            'created_at': order.created_at.strftime('%Y-%m-%d %H:%M:%S'),
            'updated_at': order.updated_at.strftime('%Y-%m-%d %H:%M:%S')
        }
        
        return JsonResponse({'success': True, 'order': order_data})
    
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=400)

@csrf_exempt
@login_required
def get_payment_status(request, order_id):
    if not request.user.is_staff:
        return JsonResponse({'error': 'Unauthorized'}, status=403)
    
    try:
        order = get_object_or_404(Order, id=order_id)
        
        payment_data = {
            'order_id': order.order_id,
            'payment_status': order.get_payment_status_display(),
            'payment_method': order.get_payment_method_display(),
            'total_amount': float(order.total_amount),
            'paid_amount': float(order.paid_amount) if hasattr(order, 'paid_amount') else 0.0
        }
        
        return JsonResponse({'success': True, 'payment': payment_data})
    
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=400)

def payment_success_view(request):
    order_id = request.GET.get('order_id')
    if not order_id:
        return render(request, 'error.html', {'error': 'Order ID not provided'})
    
    try:
        order = get_object_or_404(Order, order_id=order_id)
        context = {
            'order': order
        }
        return render(request, 'payment_success.html', context)
    except Exception as e:
        return render(request, 'error.html', {'error': 'Order not found'})

def payment_cancel_view(request):
    order_id = request.GET.get('order_id')
    if not order_id:
        return render(request, 'error.html', {'error': 'Order ID not provided'})
    
    try:
        order = get_object_or_404(Order, order_id=order_id)
        # Update order status to cancelled
        order.order_status = 'cancelled'
        order.save()
        
        context = {
            'order': order
        }
        return render(request, 'payment_cancel.html', context)
    except Exception as e:
        return render(request, 'error.html', {'error': 'Order not found'})

@csrf_exempt
def mpesa_callback(request):
    """Handle M-Pesa payment callbacks"""
    if request.method != 'POST':
        return JsonResponse({'error': 'Invalid request method'}, status=405)
    
    try:
        # Parse M-Pesa callback data
        data = json.loads(request.body)
        
        # Extract transaction details
        transaction_id = data.get('trans_id')
        order_id = data.get('order_id')
        status = data.get('status')
        
        if not all([transaction_id, order_id, status]):
            return JsonResponse({'error': 'Missing required fields'}, status=400)
        
        # Find and update order
        order = get_object_or_404(Order, order_id=order_id)
        
        if status == 'success':
            order.payment_status = 'paid'
            order.order_status = 'processing'
        else:
            order.payment_status = 'failed'
        
        order.save()
        
        return JsonResponse({'success': True, 'message': 'Callback processed successfully'})
    
    except json.JSONDecodeError:
        return JsonResponse({'error': 'Invalid JSON'}, status=400)
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=400)

@csrf_exempt
def api_search_counties(request):
    """API endpoint for searching counties"""
    if request.method != 'GET':
        return JsonResponse({'error': 'Invalid request method'}, status=405)
    
    query = request.GET.get('q', '').strip()
    
    # Sample Kenyan counties data
    counties = [
        'Nairobi', 'Mombasa', 'Kisumu', 'Nakuru', 'Eldoret', 
        'Kisii', 'Kakamega', 'Bungoma', 'Busia', 'Siaya',
        'Homa Bay', 'Migori', 'Kericho', 'Bomet', 'Narok',
        'Kajiado', 'Kitui', 'Machakos', 'Makueni', 'Meru',
        'Embu', 'Tharaka Nithi', 'Nyeri', 'Kirinyaga', 'Muranga',
        'Kiambu', 'Turkana', 'West Pokot', 'Samburu', 'Trans Nzoia',
        'Uasin Gishu', 'Elgeyo Marakwet', 'Nandi', 'Baringo', 'Laikipia',
        'Nyandarua', 'Taita Taveta', 'Kwale', 'Kilifi', 'Tana River',
        'Lamu', 'Garissa', 'Wajir', 'Mandera', 'Marsabit', 'Isiolo'
    ]
    
    if query:
        filtered_counties = [county for county in counties if query.lower() in county.lower()]
    else:
        filtered_counties = counties[:10]  # Return first 10 if no query
    
    return JsonResponse({'counties': filtered_counties})

@csrf_exempt
def api_get_cities(request):
    """API endpoint for getting cities in a county"""
    if request.method != 'GET':
        return JsonResponse({'error': 'Invalid request method'}, status=405)
    
    county = request.GET.get('county', '').strip()
    
    # Sample cities data for Kenyan counties
    cities_data = {
        'Nairobi': ['Nairobi CBD', 'Westlands', 'Karen', 'Langata', 'Kasarani', 'Embakasi', 'Dagoretti', 'Kibera', 'Ruaraka'],
        'Mombasa': ['Mombasa Island', 'Likoni', 'Changamwe', 'Kisauni', 'Nyali', 'Jomvu', 'Mvita'],
        'Kisumu': ['Kisumu CBD', 'Kondele', 'Manyatta', 'Nyalenda', 'Obunga', 'Riat'],
        'Nakuru': ['Nakuru Town', 'Bahati', 'Njoro', 'Naivasha', 'Gilgil', 'Molo'],
        'Eldoret': ['Eldoret Town', 'Kapsowar', 'Kapsabet', 'Iten', 'Cheptiret'],
        'Kiambu': ['Thika', 'Kikuyu', 'Limuru', 'Ruiru', 'Juja', 'Gigiri'],
        'Kajiado': ['Kitengela', 'Kajiado Town', 'Ngong', 'Ongata Rongai', 'Isinya'],
        'Machakos': ['Machakos Town', 'Athi River', 'Mwala', 'Kathiani', 'Yatta'],
        'Kilifi': ['Kilifi Town', 'Malindi', 'Mtwapa', 'Watamu', 'Kilifi Creek'],
        'Uasin Gishu': ['Eldoret', 'Burnt Forest', 'Moiben', 'Soy', 'Ziwa']
    }
    
    if county and county in cities_data:
        cities = cities_data[county]
    else:
        cities = ['City not found']
    
    return JsonResponse({'cities': cities})

@csrf_exempt
def api_get_delivery_info(request):
    """API endpoint for getting delivery information"""
    if request.method != 'GET':
        return JsonResponse({'error': 'Invalid request method'}, status=405)
    
    county = request.GET.get('county', '').strip()
    city = request.GET.get('city', '').strip()
    
    # Sample delivery fees and times
    delivery_info = {
        'Nairobi': {'fee': 200, 'time': '1-2 days'},
        'Mombasa': {'fee': 350, 'time': '2-3 days'},
        'Kisumu': {'fee': 250, 'time': '2-3 days'},
        'Nakuru': {'fee': 250, 'time': '2-3 days'},
        'Eldoret': {'fee': 300, 'time': '3-4 days'},
        'Kisii': {'fee': 350, 'time': '3-4 days'},
        'Kakamega': {'fee': 350, 'time': '3-4 days'},
        'Bungoma': {'fee': 400, 'time': '4-5 days'},
        'Busia': {'fee': 400, 'time': '4-5 days'},
        'Siaya': {'fee': 350, 'time': '3-4 days'},
        'Homa Bay': {'fee': 350, 'time': '3-4 days'},
        'Migori': {'fee': 350, 'time': '3-4 days'},
        'Kericho': {'fee': 300, 'time': '2-3 days'},
        'Bomet': {'fee': 300, 'time': '2-3 days'},
        'Narok': {'fee': 300, 'time': '2-3 days'},
        'Kajiado': {'fee': 250, 'time': '2-3 days'},
        'Kitui': {'fee': 400, 'time': '4-5 days'},
        'Machakos': {'fee': 250, 'time': '2-3 days'},
        'Makueni': {'fee': 300, 'time': '3-4 days'},
        'Meru': {'fee': 350, 'time': '3-4 days'},
        'Embu': {'fee': 300, 'time': '2-3 days'},
        'Tharaka Nithi': {'fee': 350, 'time': '3-4 days'},
        'Nyeri': {'fee': 300, 'time': '2-3 days'},
        'Kirinyaga': {'fee': 300, 'time': '2-3 days'},
        'Muranga': {'fee': 250, 'time': '2-3 days'},
        'Kiambu': {'fee': 200, 'time': '1-2 days'},
        'Turkana': {'fee': 500, 'time': '5-7 days'},
        'West Pokot': {'fee': 500, 'time': '5-7 days'},
        'Samburu': {'fee': 500, 'time': '5-7 days'},
        'Trans Nzoia': {'fee': 400, 'time': '4-5 days'},
        'Uasin Gishu': {'fee': 300, 'time': '3-4 days'},
        'Elgeyo Marakwet': {'fee': 400, 'time': '4-5 days'},
        'Nandi': {'fee': 350, 'time': '3-4 days'},
        'Baringo': {'fee': 450, 'time': '4-5 days'},
        'Laikipia': {'fee': 400, 'time': '4-5 days'},
        'Nyandarua': {'fee': 300, 'time': '3-4 days'},
        'Taita Taveta': {'fee': 400, 'time': '4-5 days'},
        'Kwale': {'fee': 400, 'time': '4-5 days'},
        'Kilifi': {'fee': 350, 'time': '3-4 days'},
        'Tana River': {'fee': 500, 'time': '5-7 days'},
        'Lamu': {'fee': 500, 'time': '5-7 days'},
        'Garissa': {'fee': 600, 'time': '7-10 days'},
        'Wajir': {'fee': 600, 'time': '7-10 days'},
        'Mandera': {'fee': 700, 'time': '7-10 days'},
        'Marsabit': {'fee': 600, 'time': '7-10 days'},
        'Isiolo': {'fee': 500, 'time': '5-7 days'}
    }
    
    if county in delivery_info:
        info = delivery_info[county]
    else:
        info = {'fee': 500, 'time': '5-7 days'}  # Default for unknown counties
    
    return JsonResponse({
        'delivery_fee': info['fee'],
        'delivery_time': info['time'],
        'county': county,
        'city': city
    })

@csrf_exempt
def api_calculate_delivery_fee(request):
    """API endpoint for calculating delivery fee based on location"""
    if request.method != 'GET':
        return JsonResponse({'error': 'Invalid request method'}, status=405)
    
    county = request.GET.get('county', '').strip()
    city = request.GET.get('city', '').strip()
    
    # Base delivery fees by county
    base_fees = {
        'Nairobi': 200,
        'Mombasa': 350,
        'Kisumu': 250,
        'Nakuru': 250,
        'Eldoret': 300,
        'Kisii': 350,
        'Kakamega': 350,
        'Bungoma': 400,
        'Busia': 400,
        'Siaya': 350,
        'Homa Bay': 350,
        'Migori': 350,
        'Kericho': 300,
        'Bomet': 300,
        'Narok': 300,
        'Kajiado': 250,
        'Kitui': 400,
        'Machakos': 250,
        'Makueni': 300,
        'Meru': 350,
        'Embu': 300,
        'Tharaka Nithi': 350,
        'Nyeri': 300,
        'Kirinyaga': 300,
        'Muranga': 250,
        'Kiambu': 200,
        'Turkana': 500,
        'West Pokot': 500,
        'Samburu': 500,
        'Trans Nzoia': 400,
        'Uasin Gishu': 300,
        'Elgeyo Marakwet': 400,
        'Nandi': 350,
        'Baringo': 450,
        'Laikipia': 400,
        'Nyandarua': 300,
        'Taita Taveta': 400,
        'Kwale': 400,
        'Kilifi': 350,
        'Tana River': 500,
        'Lamu': 500,
        'Garissa': 600,
        'Wajir': 600,
        'Mandera': 700,
        'Marsabit': 600,
        'Isiolo': 500
    }
    
    # Get base fee or default
    base_fee = base_fees.get(county, 500)
    
    # Additional fees for remote areas
    remote_cities = ['Kibera', 'Obunga', 'Marsabit', 'Mandera', 'Wajir', 'Garissa', 'Lamu', 'Turkana']
    if city in remote_cities:
        base_fee += 200
    
    return JsonResponse({
        'delivery_fee': base_fee,
        'county': county,
        'city': city,
        'base_fee': base_fees.get(county, 500),
        'additional_fee': 200 if city in remote_cities else 0
    })
